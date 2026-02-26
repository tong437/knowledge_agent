"""
文档处理器，支持 TXT、Markdown 和 Word 文件。
"""

import re
from pathlib import Path
from typing import Dict, Any, List

from .base_processor import BaseDataSourceProcessor
from core.models import DataSource, SourceType
from core.exceptions import ProcessingError


class DocumentProcessor(BaseDataSourceProcessor):
    """
    文档文件处理器（TXT、Markdown、Word）。

    支持：
    - 纯文本文件（.txt）
    - Markdown 文件（.md、.markdown）
    - Word 文档（.docx）- 需要 python-docx
    """

    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.docx': 'word',
    }

    def __init__(self):
        """初始化文档处理器。"""
        super().__init__()
        self._docx_available = self._check_docx_support()

    def _check_docx_support(self) -> bool:
        """检查 python-docx 是否可用。"""
        try:
            import docx
            return True
        except ImportError:
            self.logger.warning(
                "python-docx not available. Word document support disabled. "
                "Install with: pip install python-docx"
            )
            return False

    def get_supported_types(self) -> List[SourceType]:
        """获取此处理器支持的数据源类型列表。"""
        return [SourceType.DOCUMENT]

    def validate(self, source: DataSource) -> bool:
        """
        验证文档数据源是否可以被处理。

        Args:
            source: 待验证的数据源

        Returns:
            bool: 可处理返回 True，否则返回 False
        """
        if not super().validate(source):
            return False

        path = Path(source.path)
        ext = path.suffix.lower()

        if ext not in self.SUPPORTED_EXTENSIONS:
            self.logger.warning(
                f"Unsupported document extension: {ext}. "
                f"Supported: {list(self.SUPPORTED_EXTENSIONS.keys())}"
            )
            return False

        if ext == '.docx' and not self._docx_available:
            self.logger.warning(
                f"Word document support not available for: {source.path}"
            )
            return False

        return True

    def get_metadata(self, source: DataSource) -> Dict[str, Any]:
        """
        从文档数据源中提取元数据。

        Args:
            source: 待提取元数据的数据源

        Returns:
            Dict[str, Any]: 元数据字典
        """
        metadata = super().get_metadata(source)

        path = Path(source.path)
        ext = path.suffix.lower()

        metadata['document_type'] = self.SUPPORTED_EXTENSIONS.get(ext, 'unknown')

        if ext == '.docx' and self._docx_available:
            try:
                import docx
                doc = docx.Document(source.path)
                core_props = doc.core_properties

                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.title:
                    metadata['doc_title'] = core_props.title
                if core_props.subject:
                    metadata['subject'] = core_props.subject
                if core_props.created:
                    metadata['doc_created'] = core_props.created.isoformat()
                if core_props.modified:
                    metadata['doc_modified'] = core_props.modified.isoformat()

            except Exception as e:
                self.logger.warning(
                    f"Error extracting Word metadata from {source.path}: {str(e)}"
                )

        return metadata

    def _extract_content(self, source: DataSource) -> str:
        """
        从文档中提取内容。

        Args:
            source: 待提取内容的数据源

        Returns:
            str: 提取的内容

        Raises:
            ProcessingError: 内容提取失败时抛出
        """
        path = Path(source.path)
        ext = path.suffix.lower()

        if ext == '.docx':
            return self._extract_word_content(source)
        else:
            return self._extract_text_content(source)

    def _extract_text_content(self, source: DataSource) -> str:
        """
        从纯文本或 Markdown 文件中提取内容。

        Args:
            source: 数据源

        Returns:
            str: 提取的内容
        """
        encoding = source.encoding or 'utf-8'
        content = self._read_file(source.path, encoding)
        content = self._clean_text(content)
        return content

    def _extract_word_content(self, source: DataSource) -> str:
        """
        从 Word 文档中提取内容。

        Args:
            source: 数据源

        Returns:
            str: 提取的内容

        Raises:
            ProcessingError: 提取失败时抛出
        """
        if not self._docx_available:
            raise ProcessingError(
                "Word document support not available. Install python-docx."
            )

        try:
            import docx
            doc = docx.Document(source.path)

            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)

            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        paragraphs.append(row_text)

            content = '\n\n'.join(paragraphs)
            content = self._clean_text(content)
            return content

        except Exception as e:
            raise ProcessingError(
                f"Error extracting Word content from {source.path}: {str(e)}"
            ) from e

    def _clean_text(self, text: str) -> str:
        """
        清理和规范化文本内容。

        Args:
            text: 待清理的文本

        Returns:
            str: 清理后的文本
        """
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        text = text.strip()
        return text

    def _generate_title(self, source: DataSource, content: str) -> str:
        """
        为文档生成标题。

        优先从 Markdown 标题、首行非空内容、文件名中提取。

        Args:
            source: 数据源
            content: 提取的内容

        Returns:
            str: 生成的标题
        """
        path = Path(source.path)
        ext = path.suffix.lower()

        if ext in ['.md', '.markdown'] and content:
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                if line.startswith('#'):
                    title = line.lstrip('#').strip()
                    if title:
                        return title

        if content:
            lines = content.strip().split('\n')
            for line in lines:
                line = line.strip()
                if line and len(line) < 100:
                    return line

        return path.stem
