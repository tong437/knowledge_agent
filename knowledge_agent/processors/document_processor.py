"""
Document processor for TXT, Markdown, and Word files.
"""

import re
from pathlib import Path
from typing import Dict, Any, List

from .base_processor import BaseDataSourceProcessor
from ..models import DataSource, SourceType
from ..core.exceptions import ProcessingError


class DocumentProcessor(BaseDataSourceProcessor):
    """
    Processor for document files (TXT, Markdown, Word).
    
    Supports:
    - Plain text files (.txt)
    - Markdown files (.md, .markdown)
    - Word documents (.docx) - requires python-docx
    """
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.markdown': 'markdown',
        '.docx': 'word',
    }
    
    def __init__(self):
        """Initialize the document processor."""
        super().__init__()
        self._docx_available = self._check_docx_support()
    
    def _check_docx_support(self) -> bool:
        """Check if python-docx is available for Word document support."""
        try:
            import docx
            return True
        except ImportError:
            self.logger.warning(
                "python-docx not available. Word document support disabled. "
                "Install with: pip install python-docx"
            )
            return False
    
    def get_supported_types(self) -> List[SourceType]:
        """Get the list of source types this processor supports."""
        return [SourceType.DOCUMENT]
    
    def validate(self, source: DataSource) -> bool:
        """
        Validate that a document source can be processed.
        
        Args:
            source: The data source to validate
            
        Returns:
            bool: True if the source can be processed, False otherwise
        """
        # Call parent validation
        if not super().validate(source):
            return False
        
        # Check file extension
        path = Path(source.path)
        ext = path.suffix.lower()
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            self.logger.warning(
                f"Unsupported document extension: {ext}. "
                f"Supported: {list(self.SUPPORTED_EXTENSIONS.keys())}"
            )
            return False
        
        # Check if Word document support is available
        if ext == '.docx' and not self._docx_available:
            self.logger.warning(
                f"Word document support not available for: {source.path}"
            )
            return False
        
        return True
    
    def get_metadata(self, source: DataSource) -> Dict[str, Any]:
        """
        Extract metadata from a document source.
        
        Args:
            source: The data source to extract metadata from
            
        Returns:
            Dict[str, Any]: Metadata dictionary
        """
        metadata = super().get_metadata(source)
        
        # Add document-specific metadata
        path = Path(source.path)
        ext = path.suffix.lower()
        
        metadata['document_type'] = self.SUPPORTED_EXTENSIONS.get(ext, 'unknown')
        
        # For Word documents, extract additional metadata
        if ext == '.docx' and self._docx_available:
            try:
                import docx
                doc = docx.Document(source.path)
                core_props = doc.core_properties
                
                if core_props.author:
                    metadata['author'] = core_props.author
                if core_props.title:
                    metadata['doc_title'] = core_props.title
                if core_props.subject:
                    metadata['subject'] = core_props.subject
                if core_props.created:
                    metadata['doc_created'] = core_props.created.isoformat()
                if core_props.modified:
                    metadata['doc_modified'] = core_props.modified.isoformat()
                    
            except Exception as e:
                self.logger.warning(
                    f"Error extracting Word metadata from {source.path}: {str(e)}"
                )
        
        return metadata
    
    def _extract_content(self, source: DataSource) -> str:
        """
        Extract content from the document.
        
        Args:
            source: The data source to extract content from
            
        Returns:
            str: The extracted content
            
        Raises:
            ProcessingError: If content extraction fails
        """
        path = Path(source.path)
        ext = path.suffix.lower()
        
        if ext == '.docx':
            return self._extract_word_content(source)
        else:
            return self._extract_text_content(source)
    
    def _extract_text_content(self, source: DataSource) -> str:
        """
        Extract content from plain text or Markdown files.
        
        Args:
            source: The data source
            
        Returns:
            str: The extracted content
        """
        encoding = source.encoding or 'utf-8'
        content = self._read_file(source.path, encoding)
        
        # Clean up the content
        content = self._clean_text(content)
        
        return content
    
    def _extract_word_content(self, source: DataSource) -> str:
        """
        Extract content from Word documents.
        
        Args:
            source: The data source
            
        Returns:
            str: The extracted content
            
        Raises:
            ProcessingError: If extraction fails
        """
        if not self._docx_available:
            raise ProcessingError(
                "Word document support not available. Install python-docx."
            )
        
        try:
            import docx
            doc = docx.Document(source.path)
            
            # Extract all paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text:
                        paragraphs.append(row_text)
            
            content = '\n\n'.join(paragraphs)
            
            # Clean up the content
            content = self._clean_text(content)
            
            return content
            
        except Exception as e:
            raise ProcessingError(
                f"Error extracting Word content from {source.path}: {str(e)}"
            ) from e
    
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text content.
        
        Args:
            text: The text to clean
            
        Returns:
            str: Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Remove leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _generate_title(self, source: DataSource, content: str) -> str:
        """
        Generate a title for the document.
        
        Tries to extract title from:
        1. Markdown heading
        2. First non-empty line
        3. Filename
        
        Args:
            source: The data source
            content: The extracted content
            
        Returns:
            str: Generated title
        """
        path = Path(source.path)
        ext = path.suffix.lower()
        
        # For Markdown, try to extract the first heading
        if ext in ['.md', '.markdown'] and content:
            lines = content.split('\n')
            for line in lines:
                line = line.strip()
                # Check for # heading
                if line.startswith('#'):
                    title = line.lstrip('#').strip()
                    if title:
                        return title
        
        # Try first non-empty line
        if content:
            lines = content.strip().split('\n')
            for line in lines:
                line = line.strip()
                if line and len(line) < 100:
                    return line
        
        # Fall back to filename
        return path.stem
